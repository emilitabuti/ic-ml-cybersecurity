#!/usr/bin/env python3
"""Gera o relatório final de IC e o resumo estendido em DOCX e PDF."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
from hashlib import sha256
from html import escape
from pathlib import Path
import re
import subprocess

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from pypdf import PdfReader
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics
from reportlab.platypus import (
    BaseDocTemplate,
    KeepTogether,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/emili/relatorio-final/relatorio-final-emili.md"
OUT_DIR = SOURCE.parent
REPORT_DOCX = OUT_DIR / "Relatorio-Final-IC-Emili-Vieira-Tabuti.docx"
REPORT_PDF = OUT_DIR / "Relatorio-Final-IC-Emili-Vieira-Tabuti.pdf"
SUMMARY_DOCX = OUT_DIR / "Resumo-Estendido-Emili-Vieira-Tabuti.docx"
SUMMARY_PDF = OUT_DIR / "Resumo-Estendido-Emili-Vieira-Tabuti.pdf"
LOG_PATH = OUT_DIR / "log-execucao.md"

INSTITUTION_HEADER = (
    "PONTIFÍCIA UNIVERSIDADE CATÓLICA DE SÃO PAULO\n"
    "PRÓ-REITORIA DE GRADUAÇÃO\n"
    "Programa Institucional de Bolsas de Iniciação Científica — PIBIC"
)
SUMMARY_HEADER = "35º Encontro de Iniciação Científica — PUC-SP"


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 0
    rows: list[list[str]] | None = None


def parse_markdown(text: str) -> list[Block]:
    lines = text.splitlines()
    blocks: list[Block] = []
    index = 0
    while index < len(lines):
        raw = lines[index].rstrip()
        stripped = raw.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("[[") and stripped.endswith("]]"):
            blocks.append(Block("marker", stripped[2:-2]))
            index += 1
            continue
        if stripped.startswith("#"):
            match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
            if match:
                blocks.append(Block("heading", match.group(2), len(match.group(1))))
                index += 1
                continue
        if stripped.startswith("|") and index + 1 < len(lines):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].strip())
                index += 1
            rows = [
                [cell.strip() for cell in line.strip("|").split("|")]
                for line in table_lines
            ]
            if len(rows) > 1 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
                rows.pop(1)
            blocks.append(Block("table", rows=rows))
            continue
        blocks.append(Block("paragraph", stripped))
        index += 1
    return blocks


def summary_blocks(blocks: list[Block]) -> list[Block]:
    start = next(
        i for i, block in enumerate(blocks)
        if block.kind == "heading" and block.text.startswith("PARTE III")
    )
    end = next(
        i for i, block in enumerate(blocks[start:], start)
        if block.kind == "marker" and block.text == "LETTER_SECTION"
    )
    return [
        block for block in blocks[start + 1 : end]
        if not (block.kind == "marker" and block.text in {"A4_SECTION"})
    ]


def plain_text(text: str) -> str:
    return re.sub(r"[*`]", "", text).strip()


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\wÀ-ÿ][\wÀ-ÿ'–-]*\b", plain_text(text)))


def markdown_to_reportlab(text: str) -> str:
    value = escape(text)
    value = re.sub(r"`([^`]+)`", r'<font name="Courier">\1</font>', value)
    value = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", value)
    value = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<i>\1</i>", value)
    value = re.sub(
        r"(TODO:[^<\n]+)",
        r'<font color="#B00020"><b>\1</b></font>',
        value,
    )
    return value


def add_docx_inline(paragraph, text: str) -> None:
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")
    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            add_docx_run(paragraph, text[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = add_docx_run(paragraph, token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = add_docx_run(paragraph, token[1:-1])
            run.font.name = "Courier New"
        else:
            run = add_docx_run(paragraph, token[1:-1])
            run.italic = True
        position = match.end()
    if position < len(text):
        add_docx_run(paragraph, text[position:])


def add_docx_run(paragraph, text: str):
    run = paragraph.add_run(text)
    if "TODO:" in text:
        run.bold = True
        run.font.color.rgb = RGBColor(176, 0, 32)
    return run


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=60, start=80, bottom=60, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.first_child_found_in("w:tblBorders")
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "insideH", "bottom"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "666666")
        borders.append(elem)
    for edge in ("left", "right", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tbl_pr.append(borders)


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_toc_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Atualize o campo para exibir o sumário."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def configure_docx_section(section, page_size: str, summary_header: bool = False) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    if page_size == "a4":
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
    else:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(SUMMARY_HEADER if summary_header else INSTITUTION_HEADER)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9 if summary_header else 8.5)
    if not summary_header:
        run.bold = True
    footer = section.footer
    footer.is_linked_to_previous = False
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.clear()
    add_page_field(footer_paragraph)


def configure_docx_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    for level in range(1, 4):
        style = styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Cm(0)
        if level == 1:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if "Cover Title" not in styles:
        cover_title = styles.add_style("Cover Title", WD_STYLE_TYPE.PARAGRAPH)
        cover_title.font.name = "Times New Roman"
        cover_title.font.size = Pt(14)
        cover_title.font.bold = True
        cover_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cover_title.paragraph_format.space_after = Pt(18)
    if "Summary Title" not in styles:
        summary_title = styles.add_style("Summary Title", WD_STYLE_TYPE.PARAGRAPH)
        summary_title.font.name = "Times New Roman"
        summary_title.font.size = Pt(14)
        summary_title.font.bold = True
        summary_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        summary_title.paragraph_format.space_after = Pt(12)
    settings = document.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def add_docx_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, row in enumerate(rows):
        for column_index, value in enumerate(row):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, "E6E6E6")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if column_index == 0
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            add_docx_inline(paragraph, value)
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(8.5)
                if row_index == 0:
                    run.bold = True
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    set_table_borders(table)


def add_docx_signatures(document: Document) -> None:
    document.add_paragraph()
    table = document.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        cell.text = "________________________________"
    table.cell(1, 0).text = "Emili Vieira Tabuti\nAluna"
    table.cell(1, 1).text = "Prof. Dr. Daniel Couto Gatti\nOrientador"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tbl_pr.append(borders)


def build_docx(blocks: list[Block], path: Path, standalone_summary: bool = False) -> None:
    document = Document()
    configure_docx_styles(document)
    configure_docx_section(
        document.sections[0],
        "a4" if standalone_summary else "letter",
        summary_header=standalone_summary,
    )
    cover_mode = False
    summary_mode = standalone_summary
    references_mode = False
    abstract_mode = False
    affiliation_lines = 0
    for block in blocks:
        if block.kind == "marker":
            if block.text == "COVER":
                cover_mode = True
            elif block.text == "PAGEBREAK":
                document.add_page_break()
                cover_mode = False
            elif block.text == "TOC":
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.first_line_indent = Cm(0)
                add_toc_field(paragraph)
            elif block.text == "A4_SECTION":
                section = document.add_section(WD_SECTION.NEW_PAGE)
                configure_docx_section(section, "a4")
                summary_mode = True
                references_mode = False
                affiliation_lines = 0
            elif block.text == "LETTER_SECTION":
                section = document.add_section(WD_SECTION.NEW_PAGE)
                configure_docx_section(section, "letter")
                summary_mode = False
                references_mode = False
            elif block.text == "SIGNATURES":
                add_docx_signatures(document)
            continue
        if block.kind == "heading":
            references_mode = "REFERÊNCIAS" in block.text.upper()
            if cover_mode:
                paragraph = document.add_paragraph(style="Cover Title")
                if block.level == 1:
                    paragraph.paragraph_format.space_before = Pt(75)
                add_docx_inline(paragraph, block.text)
            elif standalone_summary and block.level == 2:
                paragraph = document.add_paragraph(style="Summary Title")
                add_docx_inline(paragraph, block.text)
            elif summary_mode and block.level == 2 and not standalone_summary:
                paragraph = document.add_paragraph(style="Summary Title")
                add_docx_inline(paragraph, block.text)
            else:
                level = min(block.level, 3)
                paragraph = document.add_paragraph(style=f"Heading {level}")
                add_docx_inline(paragraph, block.text)
            abstract_mode = False
            continue
        if block.kind == "table":
            add_docx_table(document, block.rows or [])
            continue
        text = block.text
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        paragraph.paragraph_format.keep_together = False
        if cover_mode:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.line_spacing = 1.15
            paragraph.paragraph_format.space_after = Pt(7)
        elif summary_mode:
            paragraph.paragraph_format.first_line_indent = Cm(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.5
            if affiliation_lines < 4:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                affiliation_lines += 1
            else:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if text.startswith("**Resumo:**"):
                paragraph.paragraph_format.line_spacing = 1
                abstract_mode = True
            elif text.startswith("**Palavras-chave:**") or text.startswith(
                "**Classificação"
            ):
                paragraph.paragraph_format.line_spacing = 1
                abstract_mode = False
            if references_mode:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.line_spacing = 1
                paragraph.paragraph_format.space_after = Pt(6)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.first_line_indent = Cm(1.25)
            paragraph.paragraph_format.line_spacing = 1.5
            if text.startswith("Fonte:"):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1
                paragraph.paragraph_format.space_after = Pt(6)
            elif references_mode:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Cm(0)
                paragraph.paragraph_format.line_spacing = 1
                paragraph.paragraph_format.space_after = Pt(6)
        add_docx_inline(paragraph, text)
    document.core_properties.title = (
        "Resumo Estendido — Emili Vieira Tabuti"
        if standalone_summary
        else "Relatório Final de Iniciação Científica — Emili Vieira Tabuti"
    )
    document.core_properties.author = "Emili Vieira Tabuti"
    document.core_properties.subject = "Iniciação Científica — PUC-SP"
    document.save(path)


def register_fonts() -> tuple[str, str, str]:
    candidates = [
        (
            "/usr/share/fonts/truetype/liberation2/LiberationSerif-Regular.ttf",
            "/usr/share/fonts/truetype/liberation2/LiberationSerif-Bold.ttf",
            "/usr/share/fonts/truetype/liberation2/LiberationSerif-Italic.ttf",
        ),
        (
            "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Italic.ttf",
        ),
    ]
    for regular, bold, italic in candidates:
        if all(Path(item).exists() for item in (regular, bold, italic)):
            pdfmetrics.registerFont(TTFont("Academic", regular))
            pdfmetrics.registerFont(TTFont("Academic-Bold", bold))
            pdfmetrics.registerFont(TTFont("Academic-Italic", italic))
            pdfmetrics.registerFontFamily(
                "Academic",
                normal="Academic",
                bold="Academic-Bold",
                italic="Academic-Italic",
                boldItalic="Academic-Bold",
            )
            return "Academic", "Academic-Bold", "Academic-Italic"
    return "Times-Roman", "Times-Bold", "Times-Italic"


FONT, FONT_BOLD, FONT_ITALIC = register_fonts()


class AcademicDocTemplate(BaseDocTemplate):
    def __init__(self, filename: str, standalone_summary: bool = False, **kwargs):
        super().__init__(filename, **kwargs)
        self.standalone_summary = standalone_summary
        self.allowSplitting = True
        letter_frame = self._make_frame(letter, 2.5 * cm, 2.5 * cm, 2.5 * cm, 2.1 * cm)
        a4_frame = self._make_frame(A4, 2.0 * cm, 2.0 * cm, 2.5 * cm, 2.5 * cm)
        self.addPageTemplates(
            [
                PageTemplate(
                    id="A4" if standalone_summary else "Letter",
                    pagesize=A4 if standalone_summary else letter,
                    frames=[a4_frame if standalone_summary else letter_frame],
                    onPage=self._draw_summary_page if standalone_summary else self._draw_letter_page,
                ),
                PageTemplate(
                    id="A4",
                    pagesize=A4,
                    frames=[a4_frame],
                    onPage=self._draw_a4_page,
                )
                if not standalone_summary
                else PageTemplate(
                    id="Letter",
                    pagesize=letter,
                    frames=[letter_frame],
                    onPage=self._draw_letter_page,
                ),
            ]
        )

    @staticmethod
    def _make_frame(page_size, left, right, top, bottom):
        from reportlab.platypus import Frame

        width, height = page_size
        return Frame(
            left,
            bottom,
            width - left - right,
            height - top - bottom,
            id=f"frame-{width}-{height}",
            leftPadding=0,
            rightPadding=0,
            topPadding=0,
            bottomPadding=0,
        )

    def _draw_header_footer(self, canvas, doc, page_size, header: str, compact=False):
        canvas.setPageSize(page_size)
        width, height = page_size
        canvas.saveState()
        canvas.setFont(FONT_BOLD if not compact else FONT, 8.2 if not compact else 9)
        y = height - 0.62 * cm
        for line in header.splitlines():
            canvas.drawCentredString(width / 2, y, line)
            y -= 0.34 * cm
        canvas.setFont(FONT, 9)
        canvas.drawCentredString(width / 2, 0.68 * cm, str(doc.page))
        canvas.restoreState()

    def _draw_letter_page(self, canvas, doc):
        self._draw_header_footer(canvas, doc, letter, INSTITUTION_HEADER)

    def _draw_a4_page(self, canvas, doc):
        self._draw_header_footer(canvas, doc, A4, INSTITUTION_HEADER)

    def _draw_summary_page(self, canvas, doc):
        self._draw_header_footer(canvas, doc, A4, SUMMARY_HEADER, compact=True)

    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph) and getattr(flowable, "_toc_level", None) is not None:
            level = flowable._toc_level
            text = flowable.getPlainText()
            key = f"h-{self.page}-{abs(hash((text, self.page)))}"
            self.canv.bookmarkPage(key)
            self.canv.addOutlineEntry(text, key, level=level, closed=False)
            self.notify("TOCEntry", (level, text, self.page, key))


def pdf_styles():
    sample = getSampleStyleSheet()
    styles = {
        "body": ParagraphStyle(
            "AcademicBody",
            parent=sample["BodyText"],
            fontName=FONT,
            fontSize=11.5,
            leading=17.25,
            alignment=TA_JUSTIFY,
            firstLineIndent=1.25 * cm,
            spaceAfter=0,
            allowWidows=0,
            allowOrphans=0,
        ),
        "summary_body": ParagraphStyle(
            "SummaryBody",
            parent=sample["BodyText"],
            fontName=FONT,
            fontSize=12,
            leading=18,
            alignment=TA_JUSTIFY,
            firstLineIndent=0,
            spaceAfter=0,
        ),
        "summary_single": ParagraphStyle(
            "SummarySingle",
            parent=sample["BodyText"],
            fontName=FONT,
            fontSize=12,
            leading=12,
            alignment=TA_JUSTIFY,
            firstLineIndent=0,
            spaceAfter=0,
        ),
        "cover1": ParagraphStyle(
            "Cover1",
            parent=sample["Title"],
            fontName=FONT_BOLD,
            fontSize=14,
            leading=18,
            alignment=TA_CENTER,
            spaceBefore=2.0 * cm,
            spaceAfter=0.8 * cm,
        ),
        "cover2": ParagraphStyle(
            "Cover2",
            parent=sample["Title"],
            fontName=FONT_BOLD,
            fontSize=14,
            leading=18,
            alignment=TA_CENTER,
            spaceAfter=1.2 * cm,
        ),
        "coverbody": ParagraphStyle(
            "CoverBody",
            parent=sample["BodyText"],
            fontName=FONT,
            fontSize=12,
            leading=15,
            alignment=TA_CENTER,
            firstLineIndent=0,
            spaceAfter=6,
        ),
        "h1": ParagraphStyle(
            "Heading1Academic",
            parent=sample["Heading1"],
            fontName=FONT_BOLD,
            fontSize=12,
            leading=15,
            alignment=TA_CENTER,
            spaceBefore=10,
            spaceAfter=8,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "Heading2Academic",
            parent=sample["Heading2"],
            fontName=FONT_BOLD,
            fontSize=12,
            leading=15,
            alignment=TA_LEFT,
            spaceBefore=10,
            spaceAfter=6,
            keepWithNext=True,
        ),
        "h3": ParagraphStyle(
            "Heading3Academic",
            parent=sample["Heading3"],
            fontName=FONT_BOLD,
            fontSize=12,
            leading=15,
            alignment=TA_LEFT,
            leftIndent=0,
            spaceBefore=10,
            spaceAfter=6,
            keepWithNext=True,
        ),
        "summary_title": ParagraphStyle(
            "SummaryTitle",
            parent=sample["Title"],
            fontName=FONT_BOLD,
            fontSize=14,
            leading=17,
            alignment=TA_CENTER,
            spaceAfter=10,
        ),
        "summary_author": ParagraphStyle(
            "SummaryAuthor",
            parent=sample["BodyText"],
            fontName=FONT,
            fontSize=12,
            leading=15,
            alignment=TA_CENTER,
            firstLineIndent=0,
            spaceAfter=0,
        ),
        "caption": ParagraphStyle(
            "Caption",
            parent=sample["BodyText"],
            fontName=FONT_BOLD,
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
            firstLineIndent=0,
            spaceBefore=6,
            spaceAfter=4,
            keepWithNext=True,
        ),
        "references": ParagraphStyle(
            "References",
            parent=sample["BodyText"],
            fontName=FONT,
            fontSize=10.5,
            leading=13,
            alignment=TA_LEFT,
            firstLineIndent=0,
            spaceAfter=7,
        ),
        "summary_references": ParagraphStyle(
            "SummaryReferences",
            parent=sample["BodyText"],
            fontName=FONT,
            fontSize=12,
            leading=14,
            alignment=TA_LEFT,
            firstLineIndent=0,
            spaceAfter=6,
        ),
    }
    return styles


def make_pdf_table(rows: list[list[str]], available_width: float) -> Table:
    converted = []
    for row_index, row in enumerate(rows):
        converted.append(
            [
                Paragraph(
                    markdown_to_reportlab(value),
                    ParagraphStyle(
                        f"cell-{row_index}-{column_index}",
                        fontName=FONT_BOLD if row_index == 0 else FONT,
                        fontSize=7.7,
                        leading=9.2,
                        alignment=TA_LEFT if column_index == 0 else TA_CENTER,
                    ),
                )
                for column_index, value in enumerate(row)
            ]
        )
    column_count = max(len(row) for row in rows)
    first_width = available_width * (0.30 if column_count >= 5 else 0.38)
    remaining = (available_width - first_width) / max(1, column_count - 1)
    widths = [first_width] + [remaining] * (column_count - 1)
    table = Table(converted, colWidths=widths, repeatRows=1, hAlign="CENTER")
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E6E6E6")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("LINEABOVE", (0, 0), (-1, 0), 0.6, colors.HexColor("#555555")),
                ("LINEBELOW", (0, 0), (-1, 0), 0.6, colors.HexColor("#555555")),
                ("LINEBELOW", (0, -1), (-1, -1), 0.6, colors.HexColor("#555555")),
            ]
        )
    )
    return table


def build_pdf(blocks: list[Block], path: Path, standalone_summary: bool = False) -> None:
    styles = pdf_styles()
    doc = AcademicDocTemplate(
        str(path),
        standalone_summary=standalone_summary,
        pagesize=A4 if standalone_summary else letter,
        title=(
            "Resumo Estendido — Emili Vieira Tabuti"
            if standalone_summary
            else "Relatório Final de Iniciação Científica — Emili Vieira Tabuti"
        ),
        author="Emili Vieira Tabuti",
    )
    story = []
    cover_mode = False
    summary_mode = standalone_summary
    references_mode = False
    affiliation_lines = 0
    for block in blocks:
        if block.kind == "marker":
            if block.text == "COVER":
                cover_mode = True
            elif block.text == "PAGEBREAK":
                story.append(PageBreak())
                cover_mode = False
            elif block.text == "TOC":
                toc = TableOfContents()
                toc.levelStyles = [
                    ParagraphStyle(
                        f"TOCLevel{level}",
                        fontName=FONT,
                        fontSize=10.5,
                        leading=14,
                        leftIndent=level * 0.5 * cm,
                        firstLineIndent=0,
                        spaceBefore=2,
                    )
                    for level in range(3)
                ]
                story.append(toc)
            elif block.text == "A4_SECTION":
                story.extend([NextPageTemplate("A4"), PageBreak()])
                summary_mode = True
                references_mode = False
                affiliation_lines = 0
            elif block.text == "LETTER_SECTION":
                story.extend([NextPageTemplate("Letter"), PageBreak()])
                summary_mode = False
                references_mode = False
            elif block.text == "SIGNATURES":
                signature_rows = [
                    [
                        Paragraph("________________________________", styles["summary_author"]),
                        Paragraph("________________________________", styles["summary_author"]),
                    ],
                    [
                        Paragraph("Emili Vieira Tabuti<br/>Aluna", styles["summary_author"]),
                        Paragraph("Prof. Dr. Daniel Couto Gatti<br/>Orientador", styles["summary_author"]),
                    ],
                ]
                signature_table = Table(signature_rows, colWidths=[8 * cm, 8 * cm])
                signature_table.setStyle(
                    TableStyle(
                        [
                            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("TOPPADDING", (0, 0), (-1, -1), 8),
                        ]
                    )
                )
                story.extend([Spacer(1, 1.5 * cm), signature_table])
            continue
        if block.kind == "heading":
            references_mode = "REFERÊNCIAS" in block.text.upper()
            if cover_mode:
                style = styles["cover1"] if block.level == 1 else styles["cover2"]
                paragraph = Paragraph(markdown_to_reportlab(block.text), style)
            elif standalone_summary and block.level == 2:
                paragraph = Paragraph(markdown_to_reportlab(block.text), styles["summary_title"])
            elif summary_mode and block.level == 2:
                paragraph = Paragraph(markdown_to_reportlab(block.text), styles["summary_title"])
            else:
                style = styles[f"h{min(block.level, 3)}"]
                paragraph = Paragraph(markdown_to_reportlab(block.text), style)
                if block.text != "SUMÁRIO" and not standalone_summary:
                    paragraph._toc_level = (
                        1 if summary_mode and block.level >= 3
                        else min(block.level, 3) - 1
                    )
            story.append(paragraph)
            continue
        if block.kind == "table":
            width = (A4[0] - 4 * cm) if summary_mode else (letter[0] - 5 * cm)
            story.extend([Spacer(1, 4), make_pdf_table(block.rows or [], width), Spacer(1, 5)])
            continue
        if cover_mode:
            style = styles["coverbody"]
        elif summary_mode:
            if affiliation_lines < 4:
                style = styles["summary_author"]
                affiliation_lines += 1
            elif references_mode:
                style = styles["summary_references"]
            elif block.text.startswith("**Resumo:**") or block.text.startswith(
                "**Palavras-chave:**"
            ) or block.text.startswith("**Classificação"):
                style = styles["summary_single"]
            else:
                style = styles["summary_body"]
        else:
            if references_mode:
                style = styles["references"]
            else:
                style = (
                    styles["caption"]
                    if block.text.startswith(("**Tabela ", "Fonte:"))
                    else styles["body"]
                )
        story.append(Paragraph(markdown_to_reportlab(block.text), style))
    doc.multiBuild(story)


def docx_integrity(path: Path) -> tuple[int, int]:
    document = Document(path)
    paragraph_count = sum(1 for paragraph in document.paragraphs if paragraph.text.strip())
    table_count = len(document.tables)
    return paragraph_count, table_count


def pdf_info(path: Path) -> tuple[int, list[tuple[float, float]]]:
    reader = PdfReader(path)
    sizes = []
    for page in reader.pages:
        sizes.append((round(float(page.mediabox.width), 1), round(float(page.mediabox.height), 1)))
    return len(reader.pages), sorted(set(sizes))


def file_hash(path: Path) -> str:
    return sha256(path.read_bytes()).hexdigest()


def source_inventory() -> list[tuple[str, str]]:
    required = [
        "docs/compartilhado/relatorios-de-ic-parcial-e-final-v3.pdf",
        "docs/compartilhado/2-INSTRUCOES-PARA-FORMATACAO-DE-RESUMO-ESTENDIDO-0508-v2.docx",
        "docs/emili/RelatorioParcial - Emili Vieira Tabuti.docx",
        "docs/emili/Plano individual - Emili Vieira Tabuti.docx",
        "docs/compartilhado/ProjetoOrientador.docx",
        "docs/compartilhado/pesquisas-ml",
        "_bmad-output/emili/planning-artifacts",
        "_bmad-output/compartilhado",
        "ml-pipeline",
    ]
    result = []
    for item in required:
        path = ROOT / item
        result.append((item, "acessível" if path.exists() else "AUSENTE"))
    return result


def make_log(summary: list[Block]) -> None:
    report_pages, report_sizes = pdf_info(REPORT_PDF)
    summary_pages, summary_sizes = pdf_info(SUMMARY_PDF)
    report_docx_paragraphs, report_docx_tables = docx_integrity(REPORT_DOCX)
    summary_docx_paragraphs, summary_docx_tables = docx_integrity(SUMMARY_DOCX)
    source_text = SOURCE.read_text(encoding="utf-8")
    summary_text = "\n".join(block.text for block in summary if block.kind in {"heading", "paragraph"})
    abstract_match = re.search(r"\*\*Resumo:\*\*(.+?)\*\*Palavras-chave:", summary_text, re.S)
    abstract_words = word_count(abstract_match.group(1)) if abstract_match else 0
    summary_words = word_count(summary_text)
    todos = sorted(set(re.findall(r"TODO:[^*\n]+", source_text)))
    pytest_result = subprocess.run(
        [str(ROOT / "ml-pipeline/.venv/bin/python"), "-m", "pytest", "tests", "-q"],
        cwd=ROOT / "ml-pipeline",
        text=True,
        capture_output=True,
        check=False,
    )
    pytest_output = (pytest_result.stdout + pytest_result.stderr).replace(
        str(ROOT), "<project-root>"
    )
    pytest_tail = "\n".join(pytest_output.splitlines()[-12:])
    lines = [
        "# Log de execução — Relatório Final de IC",
        "",
        f"- Data e hora: {datetime.now().astimezone().isoformat(timespec='seconds')}",
        f"- Fonte textual: `{SOURCE.relative_to(ROOT)}`",
        "- Codificação da fonte e do log: UTF-8",
        "",
        "## 1. Arquivos obrigatórios",
        "",
    ]
    lines.extend(f"- `{name}`: {status}" for name, status in source_inventory())
    lines.extend(
        [
            "",
            "Observação: o pedido mencionava `relatorio-parcial.pdf` e `plano-individual.docx`.",
            "No repositório, ambos foram localizados com nomes completos e em formato DOCX.",
            "",
            "## 2. Etapas executadas",
            "",
            "1. Extração do modelo institucional e do guia de resumo estendido.",
            "2. Leitura do relatório parcial, plano individual e projeto do orientador.",
            "3. Auditoria dos artefatos de planejamento, implementação, código e resultados.",
            "4. Auditoria do piloto prospectivo, seus rótulos, atributos e divisão temporal.",
            "5. Conferência dos metadados bibliográficos e DOI nos arquivos dos editores.",
            "6. Redação em português acadêmico, com voz ativa e marcação de lacunas.",
            "7. Geração do relatório e do resumo estendido em DOCX e PDF.",
            "8. Validação de integridade, paginação, tamanho de papel e contagem de palavras.",
            "9. Execução da suíte automatizada do pipeline.",
            "",
            "## 3. Validações documentais",
            "",
            f"- Relatório PDF: {report_pages} páginas; tamanhos {report_sizes}; limite de 50 páginas: {'OK' if report_pages <= 50 else 'FALHA'}.",
            f"- Resumo PDF: {summary_pages} páginas; tamanhos {summary_sizes}.",
            f"- Relatório DOCX: {report_docx_paragraphs} parágrafos; {report_docx_tables} tabelas.",
            f"- Resumo DOCX: {summary_docx_paragraphs} parágrafos; {summary_docx_tables} tabelas.",
            f"- Resumo introdutório: {abstract_words} palavras; limite de 250: {'OK' if abstract_words <= 250 else 'FALHA'}.",
            f"- Resumo estendido: {summary_words} palavras; faixa de 1.000 a 2.000: {'OK' if 1000 <= summary_words <= 2000 else 'FALHA'}.",
            "- Figuras inseridas: nenhuma.",
            "- Tabelas inseridas: somente tabelas derivadas dos CSV e relatórios JSON existentes.",
            "- Citações diretas: nenhuma.",
            "- Referências principais: padrão autor-data no relatório e padrão numérico no resumo estendido.",
            f"- Fontes: Times New Roman no DOCX; {FONT} no PDF.",
            "",
            "## 4. Controles e limitações metodológicas registrados",
            "",
            "- O pré-processamento e a seleção foram ajustados somente no treino de cada fold temporal.",
            "- Nenhuma janela cruza partição, sessão, bloco temporal ou arquivo-fonte.",
            "- O teste futuro foi lido e avaliado uma única vez depois do congelamento.",
            "- A execução final usou versões de NumPy, Pandas, PyArrow e scikit-learn diferentes das registradas no protocolo.",
            "- O executor final foi criado depois do congelamento; seu hash consta no manifesto de execução.",
            "- O rótulo corresponde ao último registro da janela, sem horizonte futuro.",
            "- O piloto prospectivo controlou o tempo, mas restaram somente cinco eventos estritos.",
            "- Validação e teste receberam um evento positivo cada no conjunto estrito.",
            "- Os horizontes de 30 e 60 segundos não conservaram negativos nas partições futuras.",
            "- O UNSW-NB15 não informa com clareza quais registros pertencem a cada evento de ataque.",
            "- A análise por tipo segmenta previsões binárias; ela não é classificação multiclasse.",
            "- A seleção foi usada nos modelos finais: top-10 na Decision Tree, top-20 na LSTM e top-30 no Random Forest.",
            "",
            "## 5. Lacunas humanas",
            "",
        ]
    )
    lines.extend(f"- {todo}" for todo in todos)
    lines.extend(
        [
            "",
            "## 6. Testes do pipeline",
            "",
            f"- Comando: `ml-pipeline/.venv/bin/python -m pytest tests -q`",
            f"- Código de saída: {pytest_result.returncode}",
            "",
            "```text",
            pytest_tail,
            "```",
            "",
            "## 7. Artefatos e hashes SHA-256",
            "",
        ]
    )
    for path in (
        REPORT_DOCX,
        REPORT_PDF,
        SUMMARY_DOCX,
        SUMMARY_PDF,
        SOURCE,
        ROOT / "ml-pipeline/reports_temporal/unsw/protocol.json",
        ROOT / "ml-pipeline/reports_temporal/unsw/final_test_metrics.json",
        ROOT / "ml-pipeline/reports_temporal/unsw/final_evaluation/execution_manifest.json",
        ROOT / "ml-pipeline/models/model_rf_temporal_v2.pkl",
    ):
        lines.append(f"- `{path.relative_to(ROOT)}`: `{file_hash(path)}`")
    LOG_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    text = SOURCE.read_text(encoding="utf-8")
    blocks = parse_markdown(text)
    summary = summary_blocks(blocks)
    build_docx(blocks, REPORT_DOCX)
    build_pdf(blocks, REPORT_PDF)
    build_docx(summary, SUMMARY_DOCX, standalone_summary=True)
    build_pdf(summary, SUMMARY_PDF, standalone_summary=True)
    make_log(summary)
    print(REPORT_DOCX)
    print(REPORT_PDF)
    print(SUMMARY_DOCX)
    print(SUMMARY_PDF)
    print(LOG_PATH)


if __name__ == "__main__":
    main()
